import pdfplumber
from docx import Document
import logging
from tqdm import tqdm
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List, Dict
import os

class DocumentProcessor:
    def __init__(
        self,
        min_chunk_size = 100,
        max_chunk_size = 1000,
        chunk_overlap = 50,
    ):
        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=max_chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def extract_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX using python-docx only"""
        try:
            doc = Document(docx_path)
            
            # Extract text from paragraphs
            paragraph_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    paragraph_text.append(paragraph.text)
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # Only add non-empty cells
                            row_text.append(cell.text.strip())
                    if row_text:  # Only add non-empty rows
                        table_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n".join(paragraph_text + table_text)
            
            self.logger.info(f"Extracted {len(all_text)} characters from DOCX")
            return all_text.strip()
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            return ""

    def extract_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        try:
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n"
            
            self.logger.info(f"Extracted {len(text)} characters from PDF using pdfplumber")
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error processing PDF {pdf_path}: {str(e)}")
            return ""

    def clean_text(self, text: str) -> str:
        if not text:
            return ""
            
        try:
            # Remove extra whitespace
            text = ' '.join(text.split())
            # Remove unnecessary line breaks while preserving paragraph structure
            text = text.replace('\n\n', '[PARA]')
            text = text.replace('\n', ' ')
            text = text.replace('[PARA]', '\n\n')
            
            self.logger.info(f"Cleaned text: {len(text)} characters")
            return text
        except Exception as e:
            self.logger.error(f"Error cleaning text: {str(e)}")
            return text

    def create_chunks(self, text: str, metadata: Dict) -> List[Dict]:
        try:
            chunks = self.text_splitter.split_text(text)
            formatted_chunks = []
            
            # Extract person name from file_name
            file_name = metadata['file_name']
            name = self.extract_person_name(file_name)
            
            for i, chunk in enumerate(chunks):
                chunk_dict = {
                    'original_file': name,  # Now storing just the person's name
                    'chunk_id': f"{name}_chunk_{i}",  # Using person's name in chunk_id
                    'content': chunk
                }
                formatted_chunks.append(chunk_dict)
                
            return formatted_chunks
        except Exception as e:
            self.logger.error(f"Error creating chunks: {str(e)}")
            return []

    def extract_person_name(self, file_name: str) -> str:
        """Extract person name from filename"""
        try:
            # Remove file extension
            name_without_extension = os.path.splitext(file_name)[0]
            
            # If there are underscores, take the part after the last underscore
            if '_' in name_without_extension:
                name = name_without_extension.split('_')[-1]
            else:
                # If no underscore, use the whole name
                name = name_without_extension
                
            return name
            
        except Exception as e:
            self.logger.error(f"Error extracting person name: {str(e)}")
            return file_name
        
    def process_document(self, file_path: str) -> List[Dict]:
        """Process a single document and return its chunks"""
        file_name = os.path.basename(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = self.extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            text = self.extract_from_docx(file_path)
        else:
            self.logger.warning(f"Unsupported file format: {file_path}")
            return []

        if not text:
            return []

        # Clean the extracted text
        cleaned_text = self.clean_text(text)
        
        # Create metadata
        metadata = {
            'file_name': file_name,
            'file_path': file_path,
            'file_type': file_extension
        }
        
        # Create and return chunks
        return self.create_chunks(cleaned_text, metadata)

    def process_folder(self, input_folder: str) -> List[Dict]:
        """Process all documents in a folder and return all chunks"""
        all_chunks = []
        
        # Get list of files
        files = [f for f in os.listdir(input_folder)
                if f.lower().endswith(('.pdf', '.docx', '.doc'))]

        self.logger.info(f"Found {len(files)} files to process")

        # Process each file
        for file_name in tqdm(files):
            file_path = os.path.join(input_folder, file_name)
            self.logger.info(f"Processing {file_name}")

            try:
                # Process document and get chunks
                document_chunks = self.process_document(file_path)
                all_chunks.extend(document_chunks)
                
                self.logger.info(f"Successfully processed {file_name} into {len(document_chunks)} chunks")

            except Exception as e:
                self.logger.error(f"Error processing {file_name}: {str(e)}")
                continue

        self.logger.info(f"Processing complete! Created {len(all_chunks)} total chunks")
        return all_chunks